"""Regression test: rebuild the generic resume and verify output matches known-good.

Compares paragraph-by-paragraph text content between the existing
Generic/Resume-MATTHEW-DRUHL.docx and a freshly built version using
the same tailoring file. This catches silent content changes from
refactors or bug fixes.

Requires:
- ~/Resume/applications/Generic/Resume-MATTHEW-DRUHL.docx (known-good)
- ~/Resume/applications/Generic/tailoring-generic.json (tailoring input)
- ~/Resume/MatthewDruhl.docx (base resume)
"""

import argparse
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "scripts"))

from resume_builder import RESUME_PATH, cmd_build

GENERIC_DIR = Path.home() / "Resume" / "applications" / "Generic"
KNOWN_GOOD = GENERIC_DIR / "Resume-MATTHEW-DRUHL.docx"
TAILORING_FILE = GENERIC_DIR / "tailoring-generic.json"

pytestmark = pytest.mark.skipif(
    not KNOWN_GOOD.exists() or not TAILORING_FILE.exists() or not RESUME_PATH.exists(),
    reason=f"Requires {KNOWN_GOOD}, {TAILORING_FILE}, and {RESUME_PATH}",
)


def _extract_text(docx_path: Path) -> list[str]:
    """Extract non-empty paragraph text from a docx file."""
    doc = Document(str(docx_path))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _extract_table_text(docx_path: Path) -> list[list[str]]:
    """Extract all table cell text from a docx file."""
    doc = Document(str(docx_path))
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)
    return tables


class TestGenericResumeRegression:
    """Rebuild the generic resume and compare against known-good output."""

    def test_paragraph_content_matches(self, tmp_path):
        """Every paragraph in the rebuilt resume should match the known-good."""
        output_dir = tmp_path / "output"
        output_dir.mkdir()

        args = argparse.Namespace(
            tailoring_file=str(TAILORING_FILE),
            output_dir=str(output_dir),
        )
        cmd_build(args)

        built_files = list(output_dir.glob("Resume-*.docx"))
        assert len(built_files) == 1, f"Expected 1 docx, got {len(built_files)}"

        known_paragraphs = _extract_text(KNOWN_GOOD)
        built_paragraphs = _extract_text(built_files[0])

        # Compare paragraph counts
        assert len(built_paragraphs) == len(known_paragraphs), (
            f"Paragraph count mismatch: known={len(known_paragraphs)}, "
            f"built={len(built_paragraphs)}\n"
            f"Known: {known_paragraphs}\n"
            f"Built: {built_paragraphs}"
        )

        # Compare each paragraph
        for i, (known, built) in enumerate(zip(known_paragraphs, built_paragraphs)):
            assert built == known, (
                f"Paragraph {i} differs:\n"
                f"  Known: {known[:120]}\n"
                f"  Built: {built[:120]}"
            )

    def test_skills_table_matches(self, tmp_path):
        """The skills table in the rebuilt resume should match the known-good."""
        output_dir = tmp_path / "output"
        output_dir.mkdir()

        args = argparse.Namespace(
            tailoring_file=str(TAILORING_FILE),
            output_dir=str(output_dir),
        )
        cmd_build(args)

        built_files = list(output_dir.glob("Resume-*.docx"))
        known_tables = _extract_table_text(KNOWN_GOOD)
        built_tables = _extract_table_text(built_files[0])

        assert len(built_tables) == len(known_tables), (
            f"Table count mismatch: known={len(known_tables)}, built={len(built_tables)}"
        )

        for t_idx, (known_table, built_table) in enumerate(zip(known_tables, built_tables)):
            assert len(built_table) == len(known_table), (
                f"Table {t_idx} row count mismatch: known={len(known_table)}, built={len(built_table)}"
            )
            for r_idx, (known_row, built_row) in enumerate(zip(known_table, built_table)):
                assert built_row == known_row, (
                    f"Table {t_idx}, row {r_idx} differs:\n"
                    f"  Known: {known_row}\n"
                    f"  Built: {built_row}"
                )

    def test_section_headers_present(self, tmp_path):
        """All expected section headers should be present in the rebuilt resume."""
        output_dir = tmp_path / "output"
        output_dir.mkdir()

        args = argparse.Namespace(
            tailoring_file=str(TAILORING_FILE),
            output_dir=str(output_dir),
        )
        cmd_build(args)

        built_files = list(output_dir.glob("Resume-*.docx"))
        built_paragraphs = _extract_text(built_files[0])

        expected_sections = [
            "Technical Skills",
            "Certifications",
            "Professional Experience",
            "Additional Relevant Experience",
            "Military Service",
            "Education",
        ]
        for section in expected_sections:
            assert section in built_paragraphs, (
                f"Missing section header: '{section}'"
            )

    def test_no_missing_roles(self, tmp_path):
        """All role titles from the tailoring file should appear in the output."""
        import json

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        tailoring = json.loads(TAILORING_FILE.read_text())

        args = argparse.Namespace(
            tailoring_file=str(TAILORING_FILE),
            output_dir=str(output_dir),
        )
        cmd_build(args)

        built_files = list(output_dir.glob("Resume-*.docx"))
        built_text = "\n".join(_extract_text(built_files[0]))

        for section_key in ("experience", "additional_experience"):
            for comp in tailoring.get(section_key, []):
                for role in comp.get("roles", []):
                    assert role["title"] in built_text, (
                        f"Missing role: '{role['title']}' at '{comp['company']}'"
                    )
