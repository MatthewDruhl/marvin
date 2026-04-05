"""
Resume Builder for MARVIN

Build tailored resumes and cover letters from structured data + template.

Usage:
  uv run --with python-docx python3 resume_builder.py view
  uv run --with python-docx python3 resume_builder.py update <subcommand> [options]
  uv run --with python-docx python3 resume_builder.py build --tailoring-file FILE --output-dir DIR
  uv run --with python-docx python3 resume_builder.py cover-letter --company NAME --job-title TITLE --body-file FILE --output-dir DIR
  uv run --with python-docx python3 resume_builder.py score --tailoring-file FILE --keywords KW1,KW2,...
  uv run --with python-docx python3 resume_builder.py auto-trim --tailoring-file FILE --output-dir DIR --keywords KW1,KW2,...
"""

from __future__ import annotations

import argparse
import copy
import json
import re
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

# Paths
SCRIPT_DIR = Path(__file__).parent
SKILL_DIR = SCRIPT_DIR.parent
DATA_FILE = Path.home() / "Resume" / "data" / "resume-data.json"
TEMPLATE_FILE = Path.home() / "Resume" / "resume-template.docx"
RESUME_PATH = Path.home() / "Resume" / "MatthewDruhl.docx"

WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Data helpers
# ---------------------------------------------------------------------------

def load_data() -> dict[str, Any]:
    """Load the master resume data file."""
    if not DATA_FILE.exists():
        print(f"Error: Data file not found at {DATA_FILE}")
        sys.exit(1)
    return json.loads(DATA_FILE.read_text())


def save_data(data: dict[str, Any]) -> None:
    """Save the master resume data file."""
    DATA_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False) + "\n")
    print(f"Data saved to {DATA_FILE}")


def load_tailoring(path: str) -> dict[str, Any]:
    """Load a tailoring file."""
    p = Path(path)
    if not p.exists():
        print(f"Error: Tailoring file not found at {p}")
        sys.exit(1)
    return json.loads(p.read_text())


# ---------------------------------------------------------------------------
# VIEW command
# ---------------------------------------------------------------------------

def cmd_view(_args: argparse.Namespace) -> None:
    """Pretty-print the resume data file."""
    data = load_data()

    header = data["header"]
    print("=" * 60)
    print(f"  {header['name']}")
    print(f"  {header.get('location', '')}  |  {header['phone']}")
    print(f"  {header['email']}  |  {header['linkedin']}")
    print("=" * 60)

    print(f"\n  TITLE: {data['title']}")
    print(f"  TAGLINE: {data['tagline']}")
    print(f"\n  SUMMARY:")
    print(f"  {data['summary'][:120]}...")

    print(f"\n  KEYWORDS ({len(data['summary_keywords'])}):")
    for kw in data["summary_keywords"]:
        print(f"    - {kw}")

    print(f"\n  SKILLS ({len(data['skills'])}):")
    for s in data["skills"]:
        cats = ", ".join(s["categories"])
        print(f"    {s['name']:20s}  [{cats}]")

    print(f"\n  CERTIFICATIONS ({len(data['certifications'])}):")
    for c in data["certifications"]:
        print(f"    {c['name']}, {c['org']} ({c['platform']}), {c['date']}")

    print("\n  EXPERIENCE:")
    for company in data["experience"]:
        print(f"\n    {company['company']}, {company['location']}")
        for role in company["roles"]:
            role_type = f", {role['type']}" if role.get("type") else ""
            print(f"      {role['title']}{role_type}  ({role['start_year']} - {role['end_year']})  [max {role['max_bullets']} bullets]")
            for b in role["bullets"]:
                tags = ", ".join(b["tags"])
                print(f"        - {b['text'][:80]}...")
                print(f"          tags: [{tags}]")

    if data.get("additional_experience"):
        print("\n  ADDITIONAL EXPERIENCE:")
        for company in data["additional_experience"]:
            print(f"\n    {company['company']}, {company['location']}")
            for role in company["roles"]:
                print(f"      {role['title']}  ({role['start_year']} - {role['end_year']})")
                for b in role["bullets"]:
                    tags = ", ".join(b["tags"])
                    print(f"        - {b['text'][:80]}...")
                    print(f"          tags: [{tags}]")

    if data.get("military"):
        mil = data["military"]
        print(f"\n  MILITARY:")
        print(f"    {mil['branch']}, {mil['location']}")
        print(f"    {mil['role']}  ({mil['start']} - {mil['end']})")
        for b in mil["bullets"]:
            print(f"      - {b['text'][:80]}...")

    print(f"\n  EDUCATION:")
    for ed in data["education"]:
        print(f"    {ed['degree']}, {ed['field']}")
        print(f"      {ed['school']}, {ed['location']}  ({ed['years']})")

    print()


# ---------------------------------------------------------------------------
# UPDATE command
# ---------------------------------------------------------------------------

def cmd_update(args: argparse.Namespace) -> None:
    """Update entries in resume-data.json."""
    data = load_data()
    sub = args.update_command

    if sub == "add-skill":
        new_skill = {
            "name": args.name,
            "categories": [c.strip() for c in args.categories.split(",")]
        }
        # Check for duplicates
        existing = [s["name"].lower() for s in data["skills"]]
        if new_skill["name"].lower() in existing:
            print(f"Skill '{args.name}' already exists.")
            return
        data["skills"].append(new_skill)
        data["skills"].sort(key=lambda s: s["name"].lower())
        save_data(data)
        print(f"Added skill: {args.name} [{args.categories}]")

    elif sub == "add-cert":
        new_cert = {
            "name": args.name,
            "org": args.org,
            "platform": args.platform,
            "date": args.date
        }
        data["certifications"].insert(0, new_cert)
        save_data(data)
        print(f"Added certification: {args.name}")

    elif sub == "add-bullet":
        tags = [t.strip() for t in args.tags.split(",")]
        new_bullet = {"text": args.text, "tags": tags}

        # Find the role
        found = False
        sections = [("experience", data.get("experience", [])),
                    ("additional_experience", data.get("additional_experience", []))]
        for _section_name, companies in sections:
            for company in companies:
                for role in company["roles"]:
                    if role["title"].lower() == args.role.lower():
                        role["bullets"].append(new_bullet)
                        found = True
                        save_data(data)
                        print(f"Added bullet to {role['title']} at {company['company']}")
                        return
        if not found:
            print(f"Error: Role '{args.role}' not found.")
            # List available roles
            for _section_name, companies in sections:
                for company in companies:
                    for role in company["roles"]:
                        print(f"  - {role['title']} ({company['company']})")

    elif sub == "edit":
        # Generic JSON path edit — for advanced use
        print("Edit mode: modify resume-data.json directly or use add-skill/add-cert/add-bullet.")
        print(f"Data file: {DATA_FILE}")

    else:
        print(f"Unknown update command: {sub}")
        sys.exit(1)


# ---------------------------------------------------------------------------
# BUILD command — docx generation
# ---------------------------------------------------------------------------

def get_elem_text(elem: etree._Element) -> str:
    """Get full text from a paragraph element, preserving tab characters."""
    parts: list[str] = []
    for child in elem.iter():
        tag = etree.QName(child.tag).localname
        if tag == "t" and child.text:
            parts.append(child.text)
        elif tag == "tab":
            parts.append("\t")
    return "".join(parts).strip()


def set_elem_text(elem: etree._Element, text: str) -> None:
    """Replace all run text in a paragraph, keeping first run's formatting."""
    runs = list(elem.iter(f"{{{WML_NS}}}r"))
    if not runs:
        return
    t_elems = list(runs[0].iter(f"{{{WML_NS}}}t"))
    if t_elems:
        t_elems[0].text = text
        t_elems[0].set(qn("xml:space"), "preserve")
        for t in t_elems[1:]:
            t.getparent().remove(t)
    for run in runs[1:]:
        run.getparent().remove(run)


def find_placeholder_paragraph(body: etree._Element, placeholder: str) -> etree._Element | None:
    """Find a paragraph containing a specific placeholder text."""
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p" and placeholder in get_elem_text(child):
            return child
    return None


def clone_paragraph_with_text(
    template_para: etree._Element,
    text: str,
    bold: bool = False,
) -> etree._Element:
    """Clone a paragraph element and set its text."""
    new_para = copy.deepcopy(template_para)
    runs = list(new_para.iter(f"{{{WML_NS}}}r"))
    if runs:
        # Clear all runs except first
        for run in runs[1:]:
            run.getparent().remove(run)
        # Set text on first run
        t_elems = list(runs[0].iter(f"{{{WML_NS}}}t"))
        if t_elems:
            t_elems[0].text = text
            t_elems[0].set(qn("xml:space"), "preserve")
            for t in t_elems[1:]:
                t.getparent().remove(t)
        # Handle bold
        rPr = runs[0].find(qn("w:rPr"))
        if bold:
            if rPr is None:
                rPr = etree.SubElement(runs[0], qn("w:rPr"))
                runs[0].insert(0, rPr)
            if rPr.find(qn("w:b")) is None:
                etree.SubElement(rPr, qn("w:b"))
        else:
            if rPr is not None:
                b_elem = rPr.find(qn("w:b"))
                if b_elem is not None:
                    rPr.remove(b_elem)
        # Remove caps/smallCaps formatting — let the data control casing
        if rPr is not None:
            for caps_tag in (qn("w:caps"), qn("w:smallCaps")):
                caps_elem = rPr.find(caps_tag)
                if caps_elem is not None:
                    rPr.remove(caps_elem)
    return new_para


def set_keep_with_next(para: etree._Element) -> None:
    """Set w:keepNext on a paragraph so Word keeps it with the following paragraph.

    Prevents 'widow headers' where a role/company header sits alone at the
    bottom of a page with all its bullets on the next page.
    """
    pPr = para.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(para, qn("w:pPr"))
        para.insert(0, pPr)
    if pPr.find(qn("w:keepNext")) is None:
        etree.SubElement(pPr, qn("w:keepNext"))


def create_role_header_paragraph(
    template_para: etree._Element,
    title: str,
    role_type: str | None,
    dates: str,
) -> etree._Element:
    """Create a role header paragraph with bold title/type and tab-separated dates.

    Mimics the original format: 'Title, Type\\tDates' with bold formatting on title/type.
    """
    new_para = copy.deepcopy(template_para)
    # Clear all existing runs
    for run in list(new_para.iter(f"{{{WML_NS}}}r")):
        run.getparent().remove(run)

    # Build the title portion
    title_text = title
    if role_type:
        title_text += f", {role_type}"

    # Create bold run for title
    r1 = etree.SubElement(new_para, qn("w:r"))
    rPr1 = etree.SubElement(r1, qn("w:rPr"))
    etree.SubElement(rPr1, qn("w:b"))
    # Copy font info from template if available
    template_runs = list(template_para.iter(f"{{{WML_NS}}}r"))
    if template_runs:
        src_rPr = template_runs[0].find(qn("w:rPr"))
        if src_rPr is not None:
            for child in src_rPr:
                child_tag = etree.QName(child.tag).localname
                if child_tag not in ("b", "caps", "smallCaps"):
                    rPr1.append(copy.deepcopy(child))
    t1 = etree.SubElement(r1, qn("w:t"))
    t1.text = title_text
    t1.set(qn("xml:space"), "preserve")

    # Create tab run
    r_tab = etree.SubElement(new_para, qn("w:r"))
    etree.SubElement(r_tab, qn("w:tab"))

    # Create bold run for dates
    r2 = etree.SubElement(new_para, qn("w:r"))
    rPr2 = etree.SubElement(r2, qn("w:rPr"))
    etree.SubElement(rPr2, qn("w:b"))
    if template_runs:
        src_rPr = template_runs[0].find(qn("w:rPr"))
        if src_rPr is not None:
            for child in src_rPr:
                child_tag = etree.QName(child.tag).localname
                if child_tag not in ("b", "caps", "smallCaps"):
                    rPr2.append(copy.deepcopy(child))
    t2 = etree.SubElement(r2, qn("w:t"))
    t2.text = dates
    t2.set(qn("xml:space"), "preserve")

    return new_para


def create_bullet_paragraph(
    template_para: etree._Element,
    text: str,
) -> etree._Element:
    """Create a bullet paragraph preserving the template's style."""
    new_para = copy.deepcopy(template_para)

    # Set text
    runs = list(new_para.iter(f"{{{WML_NS}}}r"))
    if runs:
        for run in runs[1:]:
            run.getparent().remove(run)
        t_elems = list(runs[0].iter(f"{{{WML_NS}}}t"))
        if t_elems:
            t_elems[0].text = text
            t_elems[0].set(qn("xml:space"), "preserve")
            for t in t_elems[1:]:
                t.getparent().remove(t)
        # Remove bold from bullet runs
        rPr = runs[0].find(qn("w:rPr"))
        if rPr is not None:
            b_elem = rPr.find(qn("w:b"))
            if b_elem is not None:
                rPr.remove(b_elem)
    return new_para


def create_blank_paragraph(template_para: etree._Element) -> etree._Element:
    """Create a blank paragraph."""
    new_para = copy.deepcopy(template_para)
    for run in list(new_para.iter(f"{{{WML_NS}}}r")):
        run.getparent().remove(run)
    return new_para


def build_skills_table(
    template_table: etree._Element,
    skills: list[str],
    num_cols: int = 4,
    page_width_twips: int = 10800,
) -> etree._Element:
    """Build a skills table from the template table structure."""
    new_table = copy.deepcopy(template_table)

    # Set table width to full page width
    tblPr = new_table.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is not None:
            tblW.set(qn("w:w"), str(page_width_twips))
            tblW.set(qn("w:type"), "dxa")

    # Get existing rows to use as templates
    rows = list(new_table.iter(f"{{{WML_NS}}}tr"))
    template_row = copy.deepcopy(rows[0]) if rows else None

    # Remove all existing rows
    for row in rows:
        row.getparent().remove(row)

    if template_row is None:
        return new_table

    # Check how many columns the template row has
    template_cells = list(template_row.iter(f"{{{WML_NS}}}tc"))
    template_num_cols = len(template_cells)

    # If we need more columns than the template has, add cells by cloning
    if num_cols > template_num_cols and template_cells:
        # Also need to adjust table grid if present
        tbl_grid = new_table.find(qn("w:tblGrid"))
        if tbl_grid is not None:
            grid_cols = list(tbl_grid.iter(qn("w:gridCol")))
            if grid_cols:
                # Calculate new column widths (distribute evenly)
                total_width = sum(
                    int(gc.get(qn("w:w"), "0")) for gc in grid_cols
                )
                new_col_width = total_width // num_cols
                # Remove existing grid cols
                for gc in grid_cols:
                    tbl_grid.remove(gc)
                # Add new grid cols
                for _ in range(num_cols):
                    new_gc = etree.SubElement(tbl_grid, qn("w:gridCol"))
                    new_gc.set(qn("w:w"), str(new_col_width))

        # Add extra cells to template row
        for _ in range(num_cols - template_num_cols):
            new_cell = copy.deepcopy(template_cells[-1])
            # Update cell width
            tc_pr = new_cell.find(qn("w:tcPr"))
            if tc_pr is not None:
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is not None and tbl_grid is not None:
                    tc_w.set(qn("w:w"), str(new_col_width))
            template_row.append(new_cell)

        # Also adjust widths of existing cells in template row
        if tbl_grid is not None:
            for cell in template_row.iter(f"{{{WML_NS}}}tc"):
                tc_pr = cell.find(qn("w:tcPr"))
                if tc_pr is not None:
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is not None:
                        tc_w.set(qn("w:w"), str(new_col_width))

    # Sort skills alphabetically
    sorted_skills = sorted(skills, key=str.lower)

    # Calculate rows needed
    num_rows = (len(sorted_skills) + num_cols - 1) // num_cols

    for row_idx in range(num_rows):
        new_row = copy.deepcopy(template_row)
        cells = list(new_row.iter(f"{{{WML_NS}}}tc"))
        for col_idx in range(num_cols):
            skill_idx = row_idx * num_cols + col_idx
            cell = cells[col_idx] if col_idx < len(cells) else None
            if cell is not None:
                # Set cell text
                t_elems = list(cell.iter(f"{{{WML_NS}}}t"))
                if t_elems:
                    if skill_idx < len(sorted_skills):
                        t_elems[0].text = sorted_skills[skill_idx]
                    else:
                        t_elems[0].text = ""
                    for t in t_elems[1:]:
                        t.getparent().remove(t)
        new_table.append(new_row)

    return new_table


def remove_elements_between(
    body: etree._Element,
    start_idx: int,
    end_idx: int,
) -> None:
    """Remove body elements between start_idx (exclusive) and end_idx (exclusive)."""
    children = list(body)
    to_remove = children[start_idx + 1:end_idx]
    for elem in to_remove:
        body.remove(elem)


def find_section_header_index(body: etree._Element, header_text: str) -> int | None:
    """Find the index of a section header by its text."""
    children = list(body)
    for i, child in enumerate(children):
        tag = etree.QName(child.tag).localname
        if tag == "p":
            text = get_elem_text(child).strip().lower()
            if text == header_text.lower():
                return i
    return None


def find_all_section_indices(body: etree._Element) -> list[tuple[int, str]]:
    """Find all section header indices and their text."""
    sections = [
        "Technical Skills",
        "Certifications",
        "Professional Experience",
        "Additional Relevant Experience",
        "Military Service",
        "Education",
    ]
    result = []
    children = list(body)
    for i, child in enumerate(children):
        tag = etree.QName(child.tag).localname
        if tag == "p":
            text = get_elem_text(child).strip()
            if text in sections:
                result.append((i, text))
    return result


# ---------------------------------------------------------------------------
# Page estimation helpers
# ---------------------------------------------------------------------------

# Resume page capacity: 10.5pt Calibri, single-spaced, 7.5" usable width, 10" usable height
# ~14pt line height (10.5pt + leading), 720pt usable = ~51 lines per page
LINES_PER_PAGE = 51
CHARS_PER_LINE = 95  # approximate characters per line at 10.5pt Calibri on 7.5" width
# Lines consumed by fixed header area (name, contact, blanks, title, tagline, blank)
HEADER_LINES = 8


def estimate_lines(text: str) -> int:
    """Estimate how many lines a paragraph will occupy based on character count."""
    if not text:
        return 1  # blank paragraph
    return max(1, -(-len(text) // CHARS_PER_LINE))  # ceiling division


def _insert_page_break(
    insert_after: etree._Element,
    page_two_header: etree._Element,
    company_template: etree._Element | None,
    normal_template: etree._Element | None,
    current_company_name: str | None,
) -> tuple[etree._Element, bool]:
    """Insert the Page Two header, blank line, and (continued) line.

    Returns (new_insert_after, True).
    """
    insert_after.addnext(page_two_header)
    insert_after = page_two_header

    # Blank paragraph after Page Two header
    tmpl = company_template if company_template is not None else normal_template
    if tmpl is not None:
        blank = create_blank_paragraph(tmpl)
        insert_after.addnext(blank)
        insert_after = blank

    # "(continued)" line
    if current_company_name and company_template is not None:
        continued_text = f"{current_company_name} (continued)"
        cont_para = clone_paragraph_with_text(company_template, continued_text, bold=True)
        insert_after.addnext(cont_para)
        insert_after = cont_para

    return insert_after, True


def cmd_build(args: argparse.Namespace) -> None:
    """Build a tailored resume .docx from template + tailoring file."""
    tailoring = load_tailoring(args.tailoring_file)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Determine output filename
    company = tailoring.get("experience", [{}])[0].get("company", "Company")
    company_clean = company.replace(" ", "")
    output_file = output_dir / f"Resume-MATTHEW-DRUHL-{company_clean}.docx"

    # Open the original resume as our base (preserves all formatting perfectly)
    doc = Document(str(RESUME_PATH))
    body = doc.element.body

    # --- Replace header placeholders ---
    children = list(body)

    # Paragraph 0: Name (keep as-is from original)
    # Paragraph 1: Location + Phone (keep as-is)
    # Paragraph 2: Email + LinkedIn (keep as-is)

    # Paragraph 5: Title
    title = tailoring.get("title", "")
    if title:
        set_elem_text(children[5], title)

    # Paragraph 6: Tagline — keep original unless tailoring provides one
    tagline = tailoring.get("tagline", "")
    if tagline:
        set_elem_text(children[6], tagline)

    # Paragraph 8: Summary
    summary = tailoring.get("summary", "")
    if summary:
        set_elem_text(children[8], summary)

    # --- Replace keywords (paragraphs 9-12) ---
    keywords = tailoring.get("keywords", [])
    if keywords:
        # Group keywords into lines of 3
        keyword_lines = []
        for i in range(0, len(keywords), 3):
            line = " | ".join(keywords[i:i+3])
            keyword_lines.append(line)

        # Find the keyword paragraphs (bold, contain "|", between summary and Technical Skills)
        kw_paragraphs = []
        children = list(body)
        for i, child in enumerate(children):
            if i < 9 or i > 15:
                continue
            tag = etree.QName(child.tag).localname
            if tag != "p":
                continue
            text = get_elem_text(child)
            if "|" in text:
                kw_paragraphs.append((i, child))

        # Replace first keyword paragraph, remove extras, add new ones if needed
        if kw_paragraphs:
            first_kw_para = kw_paragraphs[0][1]
            # Set first line
            set_elem_text(first_kw_para, keyword_lines[0] if keyword_lines else "")

            # Remove extra keyword paragraphs
            for _idx, para in kw_paragraphs[1:]:
                body.remove(para)

            # Add additional keyword lines after the first
            insert_after = first_kw_para
            for line in keyword_lines[1:]:
                new_para = clone_paragraph_with_text(first_kw_para, line, bold=True)
                insert_after.addnext(new_para)
                insert_after = new_para

    # --- Rebuild Technical Skills table ---
    skills = tailoring.get("skills", [])
    skills_columns = tailoring.get("skills_columns", 4)
    if skills:
        children = list(body)
        for i, child in enumerate(children):
            tag = etree.QName(child.tag).localname
            if tag == "tbl":
                new_table = build_skills_table(child, skills, num_cols=skills_columns)
                child.addnext(new_table)
                body.remove(child)
                break

    # --- Rebuild Certifications ---
    certs = tailoring.get("certifications", [])
    if certs is not None:
        children = list(body)
        sections = find_all_section_indices(body)
        cert_idx = None
        next_idx = None
        for i, (idx, name) in enumerate(sections):
            if name == "Certifications":
                cert_idx = idx
                if i + 1 < len(sections):
                    next_idx = sections[i + 1][0]
                break

        if cert_idx is not None and next_idx is not None:
            # Remove everything between Certifications header and next section
            children = list(body)
            to_remove = []
            for j in range(cert_idx + 1, min(next_idx, len(children))):
                child = children[j]
                tag = etree.QName(child.tag).localname
                if tag == "p":
                    text = get_elem_text(child).strip()
                    # Stop if we hit the next section header
                    if text and text in [s[1] for s in sections]:
                        break
                    if text:
                        to_remove.append(child)

            # Use last cert paragraph as template before removing
            cert_template = to_remove[0] if to_remove else None
            for elem in to_remove:
                body.remove(elem)

            if cert_template is not None:
                # Re-find cert header position
                children = list(body)
                for k, child in enumerate(children):
                    if get_elem_text(child).strip() == "Certifications":
                        insert_after = child
                        for cert_text in certs:
                            new_para = clone_paragraph_with_text(cert_template, cert_text, bold=True)
                            insert_after.addnext(new_para)
                            insert_after = new_para
                        break

    # --- Estimate lines consumed before experience section ---
    # This running total determines where to insert the Page Two header
    compact = tailoring.get("compact", False)
    line_count = _estimate_pre_experience_lines(tailoring)

    page_break_inserted = False

    # --- Rebuild Experience sections ---
    for section_key, section_name in [
        ("experience", "Professional Experience"),
        ("additional_experience", "Additional Relevant Experience"),
    ]:
        exp_data = tailoring.get(section_key, [])
        if not exp_data:
            continue

        children = list(body)
        sections = find_all_section_indices(body)
        section_idx = None
        next_section_idx = None
        for i, (idx, name) in enumerate(sections):
            if name == section_name:
                section_idx = idx
                if i + 1 < len(sections):
                    next_section_idx = sections[i + 1][0]
                break

        if section_idx is None:
            continue

        # Collect template paragraphs from the existing content
        children = list(body)
        # Find a bullet paragraph template (ListParagraph style)
        bullet_template = None
        company_template = None
        role_template = None
        normal_template = None

        for j in range(section_idx + 1, next_section_idx or len(children)):
            child = children[j]
            tag = etree.QName(child.tag).localname
            if tag != "p":
                continue
            text = get_elem_text(child)
            if not text:
                continue
            # Check for ListParagraph style
            pPr = child.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None and pStyle.get(qn("w:val")) == "ListParagraph":
                    if bullet_template is None:
                        bullet_template = child
                    continue
            # Check if it's a bold company line (all caps company name)
            has_bold = child.find(f".//{qn('w:b')}") is not None
            if has_bold and "\t" not in text and "Page Two" not in text and "(continued)" not in text:
                if company_template is None:
                    company_template = child
            elif has_bold and "\t" in text:
                if role_template is None:
                    role_template = child
            elif not has_bold:
                if normal_template is None:
                    normal_template = child

        # Remove all content between section header and next section
        # Preserve the Page Two header (has pageBreakBefore) for reinsertion
        children = list(body)
        to_remove = []
        page_two_header = None
        in_section = False
        for j, child in enumerate(children):
            if j == section_idx:
                in_section = True
                continue
            if in_section:
                tag = etree.QName(child.tag).localname
                text = get_elem_text(child) if tag == "p" else ""
                # Stop at next section header or at sectPr
                if tag == "sectPr":
                    break
                is_next_section = text in [s[1] for s in find_all_section_indices(body)]
                if is_next_section:
                    break
                # Preserve the Page Two header element
                if tag == "p" and child.find(f".//{qn('w:pageBreakBefore')}") is not None:
                    page_two_header = child
                to_remove.append(child)

        for elem in to_remove:
            body.remove(elem)

        # Now insert new content after section header
        children = list(body)
        insert_after = None
        for child in children:
            if get_elem_text(child).strip() == section_name:
                insert_after = child
                break

        if insert_after is None:
            continue

        # Use the templates we found (fall back to normal_template or company_template)
        if bullet_template is None and normal_template is not None:
            bullet_template = normal_template
        if role_template is None and company_template is not None:
            role_template = company_template

        # Add blank paragraph after section header (skip in compact mode)
        compact = tailoring.get("compact", False)
        if not compact and (normal_template is not None or company_template is not None):
            tmpl = company_template if company_template is not None else normal_template
            blank = create_blank_paragraph(tmpl)
            insert_after.addnext(blank)
            insert_after = blank

        current_company_name = None

        for comp in exp_data:
            # Company header: "COMPANY, Location"
            company_text = f"{comp['company']}, {comp['location']}"
            current_company_name = comp["company"]
            if company_template is not None:
                para = clone_paragraph_with_text(company_template, company_text, bold=True)
                set_keep_with_next(para)
                insert_after.addnext(para)
                insert_after = para
                line_count += 1

            for role in comp.get("roles", []):
                # Check if adding the role header + at least one bullet would
                # overflow page 1.  If the header fits but the first bullet
                # doesn't, push the header to page 2 to avoid widow headers.
                role_header_lines = 1
                first_bullet = role.get("bullets", [""])[0] if role.get("bullets") else ""
                first_bullet_lines = estimate_lines(first_bullet) if first_bullet else 1
                lines_needed = role_header_lines + first_bullet_lines
                if not page_break_inserted and (line_count + lines_needed) > LINES_PER_PAGE:
                    # Insert page break before this role header
                    if page_two_header is not None:
                        insert_after, page_break_inserted = _insert_page_break(
                            insert_after, page_two_header, company_template, normal_template,
                            current_company_name,
                        )

                # Role header with dates
                dates = role.get("dates", "")
                role_type = role.get("type")
                if role_template is not None:
                    para = create_role_header_paragraph(role_template, role["title"], role_type, dates)
                    set_keep_with_next(para)
                    insert_after.addnext(para)
                    insert_after = para
                    line_count += 1

                # Bullets
                for bullet_text in role.get("bullets", []):
                    bullet_lines = estimate_lines(bullet_text)

                    # Check if adding this bullet would overflow page 1
                    if not page_break_inserted and (line_count + bullet_lines) > LINES_PER_PAGE:
                        # Insert page break before this bullet
                        if page_two_header is not None:
                            insert_after, page_break_inserted = _insert_page_break(
                                insert_after, page_two_header, company_template, normal_template,
                                current_company_name,
                            )

                    if bullet_template is not None:
                        para = create_bullet_paragraph(bullet_template, bullet_text)
                    elif normal_template is not None:
                        para = clone_paragraph_with_text(normal_template, bullet_text, bold=False)
                    else:
                        continue
                    insert_after.addnext(para)
                    insert_after = para
                    line_count += bullet_lines

        # Add trailing blank paragraph after section content (before next section header)
        if insert_after is not None:
            tmpl = company_template if company_template is not None else normal_template
            if tmpl is not None:
                blank = create_blank_paragraph(tmpl)
                insert_after.addnext(blank)

    # --- Rebuild Military section ---
    mil_data = tailoring.get("military")
    if mil_data:
        children = list(body)
        sections = find_all_section_indices(body)
        mil_idx = None
        next_idx = None
        for i, (idx, name) in enumerate(sections):
            if name == "Military Service":
                mil_idx = idx
                if i + 1 < len(sections):
                    next_idx = sections[i + 1][0]
                break

        if mil_idx is not None:
            # Find templates before removing
            children = list(body)
            mil_bullet_template = None
            mil_company_template = None
            mil_role_template = None
            for j in range(mil_idx + 1, next_idx or len(children)):
                child = children[j]
                tag = etree.QName(child.tag).localname
                if tag != "p":
                    continue
                text = get_elem_text(child)
                if not text:
                    continue
                pPr = child.find(qn("w:pPr"))
                if pPr is not None:
                    pStyle = pPr.find(qn("w:pStyle"))
                    if pStyle is not None and pStyle.get(qn("w:val")) == "ListParagraph":
                        if mil_bullet_template is None:
                            mil_bullet_template = child
                        continue
                has_bold = child.find(f".//{qn('w:b')}") is not None
                if has_bold and "\t" not in text:
                    if mil_company_template is None:
                        mil_company_template = child
                elif has_bold and "\t" in text:
                    if mil_role_template is None:
                        mil_role_template = child

            # Remove content
            children = list(body)
            to_remove = []
            in_section = False
            for j, child in enumerate(children):
                if j == mil_idx:
                    # Need to re-find since indices may have shifted
                    if get_elem_text(child).strip() == "Military Service":
                        in_section = True
                        continue
                if get_elem_text(child).strip() == "Military Service":
                    in_section = True
                    continue
                if in_section:
                    tag = etree.QName(child.tag).localname
                    if tag == "sectPr":
                        break
                    text = get_elem_text(child) if tag == "p" else ""
                    is_next = text in [s[1] for s in find_all_section_indices(body)]
                    if is_next:
                        break
                    to_remove.append(child)

            for elem in to_remove:
                body.remove(elem)

            # Re-find military header
            children = list(body)
            insert_after = None
            for child in children:
                if get_elem_text(child).strip() == "Military Service":
                    insert_after = child
                    break

            if insert_after is not None and mil_company_template is not None:
                blank = create_blank_paragraph(mil_company_template)
                insert_after.addnext(blank)
                insert_after = blank

                # Branch + location
                branch_text = f"{mil_data['branch']}, {mil_data['location']}"
                para = clone_paragraph_with_text(mil_company_template, branch_text, bold=True)
                set_keep_with_next(para)
                insert_after.addnext(para)
                insert_after = para

                # Role + dates — use mil_role_template or fall back to mil_company_template
                role_tmpl = mil_role_template if mil_role_template is not None else mil_company_template
                para = create_role_header_paragraph(
                    role_tmpl,
                    mil_data["role"],
                    None,
                    f"{mil_data['start']} \u2013 {mil_data['end']}"
                )
                set_keep_with_next(para)
                insert_after.addnext(para)
                insert_after = para

                # Bullets
                for bullet_text in mil_data.get("bullets", []):
                    if mil_bullet_template is not None:
                        para = create_bullet_paragraph(mil_bullet_template, bullet_text)
                    else:
                        para = clone_paragraph_with_text(mil_company_template, bullet_text, bold=False)
                    insert_after.addnext(para)
                    insert_after = para

                # Trailing blank paragraph
                blank = create_blank_paragraph(mil_company_template)
                insert_after.addnext(blank)

    # --- Rebuild Education ---
    edu_data = tailoring.get("education", [])
    if edu_data:
        children = list(body)
        edu_idx = None
        for i, child in enumerate(children):
            if get_elem_text(child).strip() == "Education":
                edu_idx = i
                break

        if edu_idx is not None:
            # Find template
            children = list(body)
            edu_template = None
            for j in range(edu_idx + 1, len(children)):
                child = children[j]
                tag = etree.QName(child.tag).localname
                if tag == "sectPr":
                    break
                if tag == "p" and get_elem_text(child).strip():
                    edu_template = child
                    break

            # Remove content after Education header
            children = list(body)
            to_remove = []
            in_section = False
            for child in children:
                if get_elem_text(child).strip() == "Education":
                    in_section = True
                    continue
                if in_section:
                    tag = etree.QName(child.tag).localname
                    if tag == "sectPr":
                        break
                    to_remove.append(child)

            for elem in to_remove:
                body.remove(elem)

            # Re-find Education header
            children = list(body)
            insert_after = None
            for child in children:
                if get_elem_text(child).strip() == "Education":
                    insert_after = child
                    break

            if insert_after is not None and edu_template is not None:
                blank = create_blank_paragraph(edu_template)
                insert_after.addnext(blank)
                insert_after = blank

                for ed in edu_data:
                    ed_text = f"{ed['degree']}, {ed['field']}, {ed['school']}, {ed['location']}"
                    dates = ed.get("years", "")
                    para = create_role_header_paragraph(edu_template, ed_text, None, dates)
                    insert_after.addnext(para)
                    insert_after = para

    # --- Remove "Additional role held:" lines ---
    children = list(body)
    for child in children:
        tag = etree.QName(child.tag).localname
        if tag == "p" and get_elem_text(child).strip() == "Additional role held:":
            body.remove(child)

    # --- Compact spacing: remove unnecessary blank paragraphs ---
    if tailoring.get("compact", False):
        # Section headers that should NOT have a blank paragraph after them
        section_headers = {
            "Technical Skills", "Certifications", "Professional Experience",
            "Additional Relevant Experience", "Military Service", "Education",
        }
        children = list(body)
        to_remove = []
        for i, child in enumerate(children):
            tag = etree.QName(child.tag).localname
            if tag != "p":
                continue
            text = get_elem_text(child).strip()
            if text:
                continue
            # This is a blank paragraph. Check context.
            # Remove blank paragraphs that appear:
            # 1. Right after a section header
            # 2. Between paragraphs 3-4 (the blanks before the title)
            # But keep blank before "Professional Experience" (visual separator after certs)
            prev_text = ""
            if i > 0:
                prev = children[i - 1]
                if etree.QName(prev.tag).localname == "p":
                    prev_text = get_elem_text(prev).strip()
            # Remove if previous is a section header
            if prev_text in section_headers:
                to_remove.append(child)
                continue
            # Remove if next element is a section header (blank before section)
            if i + 1 < len(children):
                next_child = children[i + 1]
                if etree.QName(next_child.tag).localname == "p":
                    next_text = get_elem_text(next_child).strip()
                    if next_text in section_headers:
                        to_remove.append(child)
                        continue
                # Remove blank before a table (skills table)
                if etree.QName(next_child.tag).localname == "tbl":
                    to_remove.append(child)
                    continue
            # Remove consecutive blanks (keep at most 1)
            elif i > 0 and children[i - 1] in to_remove:
                # Previous was already marked for removal, but this might be
                # a second blank - remove it too
                pass
            # Remove consecutive blanks in header area (keep at most 1)
            elif i < 10 and not text:
                if i > 0:
                    prev_child = children[i - 1]
                    if etree.QName(prev_child.tag).localname == "p":
                        pprev_text = get_elem_text(prev_child).strip()
                        if not pprev_text:
                            # Previous was also blank - remove this one
                            to_remove.append(child)

        for elem in to_remove:
            body.remove(elem)

    # Save
    doc.save(str(output_file))
    print(f"Resume built: {output_file}")


# ---------------------------------------------------------------------------
# COVER-LETTER command
# ---------------------------------------------------------------------------

def cmd_cover_letter(args: argparse.Namespace) -> None:
    """Build a cover letter .docx."""
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    body_file = Path(args.body_file)
    if not body_file.exists():
        print(f"Error: Body file not found: {body_file}")
        sys.exit(1)

    body_paragraphs = body_file.read_text().strip().split("\n\n")
    company = args.company
    job_title = args.job_title
    date_str = args.date or datetime.now().strftime("%B %d, %Y")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Date and company
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{date_str}\n{company}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Salutation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Dear Hiring Manager,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Body paragraphs
    for para_text in body_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(para_text.strip())
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Signature block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Matthew Druhl")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact_lines = [
        "Empower Teams & Systems | Deliver Scalable Solutions & Seamless Collaboration",
        "+1 603-377-1038",
        "matthewdruhl@gmail.com",
        "www.linkedin.com/in/matthew-druhl",
    ]
    for i, line in enumerate(contact_lines):
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(contact_lines) - 1:
            p.add_run("\n").font.size = Pt(11)

    # Save
    company_clean = company.replace(" ", "")
    output_file = output_dir / f"CoverLetter-MATTHEW-DRUHL-{company_clean}.docx"
    doc.save(str(output_file))
    print(f"Cover letter created: {output_file}")


# ---------------------------------------------------------------------------
# SCORE command — confidence scoring for bullets
# ---------------------------------------------------------------------------

def score_bullet(bullet_text: str, keywords: list[str]) -> float:
    """Score a bullet against job posting keywords. Returns 0.0-1.0.

    Uses word-boundary matching so 'SQL' won't match 'MySQL' and
    'Java' won't match 'JavaScript'.
    """
    if not keywords:
        return 0.5
    matches = sum(
        1 for kw in keywords
        if re.search(r'\b' + re.escape(kw) + r'\b', bullet_text, re.IGNORECASE)
    )
    return matches / len(keywords)


def score_tailoring(tailoring: dict[str, Any], keywords: list[str]) -> list[dict[str, Any]]:
    """Score all bullets in a tailoring file against keywords.

    Returns a list of dicts with company, role, bullet_index, text, score,
    sorted by score ascending (lowest first = first to cut).
    """
    scored: list[dict[str, Any]] = []

    for section_key in ("experience", "additional_experience"):
        for comp in tailoring.get(section_key, []):
            for role in comp.get("roles", []):
                for i, bullet_text in enumerate(role.get("bullets", [])):
                    score = score_bullet(bullet_text, keywords)
                    scored.append({
                        "section": section_key,
                        "company": comp["company"],
                        "role": role["title"],
                        "bullet_index": i,
                        "text": bullet_text[:80] + ("..." if len(bullet_text) > 80 else ""),
                        "full_text": bullet_text,
                        "score": score,
                    })

    # Military bullets scored lower priority (usually kept)
    mil = tailoring.get("military")
    if mil:
        for i, bullet_text in enumerate(mil.get("bullets", [])):
            score = score_bullet(bullet_text, keywords)
            scored.append({
                "section": "military",
                "company": mil.get("branch", "Military"),
                "role": mil.get("role", ""),
                "bullet_index": i,
                "text": bullet_text[:80] + ("..." if len(bullet_text) > 80 else ""),
                "full_text": bullet_text,
                "score": score,
            })

    scored.sort(key=lambda x: x["score"])
    return scored


def cmd_score(args: argparse.Namespace) -> None:
    """Score bullets in a tailoring file against job keywords."""
    tailoring = load_tailoring(args.tailoring_file)
    keywords = [kw.strip() for kw in args.keywords.split(",")]
    scored = score_tailoring(tailoring, keywords)

    print(f"\n{'=' * 70}")
    print(f"  Bullet Scoring — {len(scored)} bullets vs {len(keywords)} keywords")
    print(f"  Keywords: {', '.join(keywords)}")
    print(f"{'=' * 70}\n")

    for entry in scored:
        bar = "#" * int(entry["score"] * 20)
        print(f"  {entry['score']:.2f} [{bar:20s}] {entry['company']} / {entry['role']}")
        print(f"       {entry['text']}")
        print()


# ---------------------------------------------------------------------------
# AUTO-TRIM command — iterative build + trim to 2 pages
# ---------------------------------------------------------------------------

def remove_lowest_bullet(tailoring: dict[str, Any], scored: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Remove the lowest-scored bullet from the tailoring data.

    Skips roles that only have 1 bullet (never leave a role with 0 bullets).
    Returns the removed entry, or None if nothing can be removed.
    """
    for entry in scored:
        section_key = entry["section"]

        if section_key == "military":
            mil = tailoring.get("military", {})
            bullets = mil.get("bullets", [])
            if len(bullets) <= 1:
                continue
            if entry["bullet_index"] < len(bullets):
                bullets.pop(entry["bullet_index"])
                return entry
            continue

        for comp in tailoring.get(section_key, []):
            if comp["company"] != entry["company"]:
                continue
            for role in comp.get("roles", []):
                if role["title"] != entry["role"]:
                    continue
                bullets = role.get("bullets", [])
                if len(bullets) <= 1:
                    continue
                if entry["bullet_index"] < len(bullets):
                    bullets.pop(entry["bullet_index"])
                    return entry

    return None


def _estimate_pre_experience_lines(tailoring: dict[str, Any]) -> int:
    """Estimate lines consumed before the experience section.

    Shared by cmd_build (for page break placement) and estimate_total_lines
    (for auto-trim). Single source of truth for this calculation.
    """
    compact = tailoring.get("compact", False)
    lines = HEADER_LINES if not compact else HEADER_LINES - 1

    summary = tailoring.get("summary", "")
    if summary:
        lines += estimate_lines(summary)

    keywords = tailoring.get("keywords", [])
    lines += -(-len(keywords) // 3) if keywords else 0
    if not compact:
        lines += 1  # blank after keywords

    lines += 1  # Technical Skills header
    if not compact:
        lines += 1  # blank after header
    skills = tailoring.get("skills", [])
    skills_cols = tailoring.get("skills_columns", 4)
    lines += -(-len(skills) // skills_cols) + 1 if skills else 0
    if not compact:
        lines += 1  # blank after table

    lines += 1  # Certifications header
    certs = tailoring.get("certifications", [])
    lines += len(certs)

    lines += 1  # Prof Experience header
    if not compact:
        lines += 1  # blank after header

    return lines


def estimate_total_lines(tailoring: dict[str, Any]) -> int:
    """Estimate total lines a tailoring file will produce in the resume."""
    compact = tailoring.get("compact", False)
    lines = _estimate_pre_experience_lines(tailoring)

    # Experience sections (minus the first Prof Experience header already counted)
    first_section = True
    for section_key in ("experience", "additional_experience"):
        exp_data = tailoring.get(section_key, [])
        if not exp_data:
            continue
        if not first_section:
            lines += 1  # section header
            if not compact:
                lines += 1  # blank after header
        first_section = False
        for comp in exp_data:
            lines += 1  # company header
            for role in comp.get("roles", []):
                lines += 1  # role header
                for bullet_text in role.get("bullets", []):
                    lines += estimate_lines(bullet_text)
        lines += 1  # trailing blank

    # Military
    mil = tailoring.get("military")
    if mil:
        lines += 1  # header
        if not compact:
            lines += 1
        lines += 2  # branch + role headers
        for bullet_text in mil.get("bullets", []):
            lines += estimate_lines(bullet_text)
        lines += 1  # trailing blank

    # Education
    edu = tailoring.get("education", [])
    if edu:
        lines += 1  # header
        if not compact:
            lines += 1
        lines += len(edu)

    return lines


def estimate_pages(tailoring: dict[str, Any]) -> int:
    """Estimate how many pages a tailoring file will produce."""
    total_lines = estimate_total_lines(tailoring)
    # Page 1 has full capacity, page 2 loses ~5 lines for "Page Two" header + continued
    if total_lines <= LINES_PER_PAGE:
        return 1
    remaining = total_lines - LINES_PER_PAGE
    page2_capacity = LINES_PER_PAGE - 5  # account for page break overhead
    extra_pages = -(-remaining // page2_capacity)  # ceiling division
    return 1 + extra_pages


def cmd_auto_trim(args: argparse.Namespace) -> None:
    """Iteratively trim lowest-scored bullets until resume fits 2 pages.

    Uses fast line estimation for the trim loop, then builds the final
    DOCX and verifies with Word conversion (if available).
    """
    tailoring = load_tailoring(args.tailoring_file)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    keywords = [kw.strip() for kw in args.keywords.split(",")]
    max_pages = args.max_pages

    removed: list[dict[str, Any]] = []
    max_iterations = 20

    # Phase 1: Fast estimation loop
    print(f"\n  Phase 1: Estimating and trimming...")
    total_lines = estimate_total_lines(tailoring)
    est_pages = estimate_pages(tailoring)
    total_bullets = sum(
        len(role.get("bullets", []))
        for section in ("experience", "additional_experience")
        for comp in tailoring.get(section, [])
        for role in comp.get("roles", [])
    )
    mil_bullets = len(tailoring.get("military", {}).get("bullets", []))
    total_bullets += mil_bullets
    print(f"  Initial: ~{total_lines} lines, ~{est_pages} pages, {total_bullets} bullets")

    iteration = 0
    while est_pages > max_pages and iteration < max_iterations:
        iteration += 1
        scored = score_tailoring(tailoring, keywords)
        entry = remove_lowest_bullet(tailoring, scored)
        if entry is None:
            print("  Cannot remove more bullets (all roles at 1 bullet minimum).")
            break

        removed.append(entry)
        total_lines = estimate_total_lines(tailoring)
        est_pages = estimate_pages(tailoring)
        print(f"    Removed [{entry['score']:.2f}]: {entry['company']} / {entry['role']}")
        print(f"             {entry['text']}")
        print(f"             ~{total_lines} lines, ~{est_pages} pages")

    # Phase 2: Build the final DOCX
    print(f"\n  Phase 2: Building final resume...")
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".json", delete=False, dir=output_dir
    ) as f:
        json.dump(tailoring, f, indent=2)
        temp_tailoring_path = Path(f.name)

    build_args = argparse.Namespace(
        tailoring_file=str(temp_tailoring_path),
        output_dir=str(output_dir),
    )
    cmd_build(build_args)
    temp_tailoring_path.unlink(missing_ok=True)

    company = tailoring.get("experience", [{}])[0].get("company", "Company")
    company_clean = company.replace(" ", "")
    output_file = output_dir / f"Resume-MATTHEW-DRUHL-{company_clean}.docx"

    # Rename output file if --company provided
    company_name = getattr(args, "company", "") or ""
    if company_name:
        # Find the built file and rename it
        docx_files = sorted(
            output_dir.glob("Resume-MATTHEW-DRUHL-*.docx"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
        if docx_files:
            output_file = docx_files[0]
            new_name = output_dir / f"Resume-MATTHEW-DRUHL-{company_name.upper().replace(' ', '')}.docx"
            if output_file != new_name:
                output_file.rename(new_name)
                output_file = new_name
    else:
        docx_files = sorted(
            output_dir.glob("Resume-MATTHEW-DRUHL-*.docx"),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
        if docx_files:
            output_file = docx_files[0]

    total_bullets = sum(
        len(role.get("bullets", []))
        for section in ("experience", "additional_experience")
        for comp in tailoring.get(section, [])
        for role in comp.get("roles", [])
    )
    mil_bullets = len(tailoring.get("military", {}).get("bullets", []))
    total_bullets += mil_bullets

    # Final summary
    print(f"\n{'=' * 60}")
    print(f"  Auto-Trim Complete")
    print(f"  Final: ~{est_pages} pages (estimated), {total_bullets} bullets")
    print(f"  Output: {output_file}")
    print(f"{'=' * 60}")

    if removed:
        print(f"\n  Bullets removed ({len(removed)} total, lowest relevance first):")
        for entry in removed:
            print(f"    [{entry['score']:.2f}] {entry['company']} / {entry['role']}")
            print(f"          {entry['text']}")

    # Save final tailoring file
    final_tailoring_path = output_dir / "tailoring-trimmed.json"
    final_tailoring_path.write_text(json.dumps(tailoring, indent=2) + "\n")
    print(f"\n  Trimmed tailoring saved: {final_tailoring_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Resume Builder for MARVIN",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    # view
    subparsers.add_parser("view", help="Pretty-print resume data")

    # update
    update_parser = subparsers.add_parser("update", help="Update resume data")
    update_sub = update_parser.add_subparsers(dest="update_command", help="Update subcommands")

    # update add-skill
    add_skill = update_sub.add_parser("add-skill", help="Add a skill")
    add_skill.add_argument("--name", required=True, help="Skill name")
    add_skill.add_argument("--categories", required=True, help="Comma-separated categories")

    # update add-cert
    add_cert = update_sub.add_parser("add-cert", help="Add a certification")
    add_cert.add_argument("--name", required=True, help="Certification name")
    add_cert.add_argument("--org", required=True, help="Issuing organization")
    add_cert.add_argument("--platform", required=True, help="Platform (e.g., Coursera)")
    add_cert.add_argument("--date", required=True, help="Date earned")

    # update add-bullet
    add_bullet = update_sub.add_parser("add-bullet", help="Add a bullet to a role")
    add_bullet.add_argument("--role", required=True, help="Role title to add bullet to")
    add_bullet.add_argument("--text", required=True, help="Bullet text")
    add_bullet.add_argument("--tags", required=True, help="Comma-separated tags")

    # update edit
    update_sub.add_parser("edit", help="Open data file for manual editing")

    # build
    build_parser = subparsers.add_parser("build", help="Build tailored resume from template + tailoring file")
    build_parser.add_argument("--tailoring-file", required=True, help="Path to tailoring JSON file")
    build_parser.add_argument("--output-dir", required=True, help="Output directory")

    # cover-letter
    cl_parser = subparsers.add_parser("cover-letter", help="Build a cover letter")
    cl_parser.add_argument("--company", required=True, help="Company name")
    cl_parser.add_argument("--job-title", required=True, help="Job title")
    cl_parser.add_argument("--body-file", required=True, help="Path to .txt file with letter body")
    cl_parser.add_argument("--output-dir", required=True, help="Output directory")
    cl_parser.add_argument("--date", default="", help="Date string (defaults to today)")

    # score
    score_parser = subparsers.add_parser("score", help="Score bullets against job keywords")
    score_parser.add_argument("--tailoring-file", required=True, help="Path to tailoring JSON file")
    score_parser.add_argument("--keywords", required=True, help="Comma-separated job keywords")

    # auto-trim
    trim_parser = subparsers.add_parser("auto-trim", help="Iteratively trim to fit page limit")
    trim_parser.add_argument("--tailoring-file", required=True, help="Path to tailoring JSON file")
    trim_parser.add_argument("--output-dir", required=True, help="Output directory")
    trim_parser.add_argument("--keywords", required=True, help="Comma-separated job keywords")
    trim_parser.add_argument("--max-pages", type=int, default=2, help="Max pages (default: 2)")
    trim_parser.add_argument("--company", default="", help="Company name for output filename (overrides tailoring)")

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    commands = {
        "view": cmd_view,
        "update": cmd_update,
        "build": cmd_build,
        "cover-letter": cmd_cover_letter,
        "score": cmd_score,
        "auto-trim": cmd_auto_trim,
    }

    commands[args.command](args)


if __name__ == "__main__":
    main()
