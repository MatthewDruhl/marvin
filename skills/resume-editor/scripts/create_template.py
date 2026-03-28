"""Create resume template by copying the original and replacing content with placeholders.

Run: uv run --with python-docx python3 create_template.py
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from pathlib import Path

RESUME_PATH = Path.home() / "Resume" / "MatthewDruhl.docx"
TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "resume-template.docx"


def get_text(elem: etree._Element) -> str:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return "".join(t.text for t in elem.iter(f"{{{ns}}}t") if t.text).strip()


def set_paragraph_text(elem: etree._Element, text: str) -> None:
    """Replace all run text with a single value, keeping first run's formatting."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    runs = list(elem.iter(f"{{{ns}}}r"))
    if not runs:
        return
    t_elems = list(runs[0].iter(f"{{{ns}}}t"))
    if t_elems:
        t_elems[0].text = text
        for t in t_elems[1:]:
            t.getparent().remove(t)
    for run in runs[1:]:
        run.getparent().remove(run)


def main() -> None:
    doc = Document(str(RESUME_PATH))
    body = doc.element.body
    children = list(body)

    # Replace header/title/summary paragraphs with placeholders
    replacements = {
        0: "{{NAME}}",
        1: "{{LOCATION_PHONE}}",
        2: "{{EMAIL_LINKEDIN}}",
        5: "{{TITLE}}",
        6: "{{TAGLINE}}",
        8: "{{SUMMARY}}",
    }

    for i, child in enumerate(children):
        tag = etree.QName(child.tag).localname
        if i in replacements and tag == "p":
            set_paragraph_text(child, replacements[i])

    # Replace keyword lines (9-12) with single placeholder, remove extras
    keyword_indices = [9, 10, 11, 12]
    for idx in keyword_indices:
        child = children[idx]
        if idx == 9:
            set_paragraph_text(child, "{{KEYWORDS}}")
        else:
            body.remove(child)

    # Save template
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(TEMPLATE_PATH))
    print(f"Template created: {TEMPLATE_PATH}")

    # Verify
    doc2 = Document(str(TEMPLATE_PATH))
    for p in doc2.paragraphs:
        t = p.text.strip()
        if t:
            print(f"  {t[:80]}")
    for table in doc2.tables:
        print("  [TABLE]")


if __name__ == "__main__":
    main()
