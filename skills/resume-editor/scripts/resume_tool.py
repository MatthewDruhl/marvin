"""
Resume Editor Tool for MARVIN

Reads and edits a Word (.docx) resume while preserving formatting.
Usage: uv run --with python-docx --with docx2pdf --with PyPDF2 python3 resume_tool.py <command> [options]

Commands:
  read               - Display resume structure and content
  page-count         - Check page count (converts to PDF via Word)
  backup             - Create a timestamped backup
  add-section        - Add a new section header
  add-entry          - Add an entry to an existing section
  add-skill          - Add a skill to the Technical Skills table
  remove-skill       - Remove a skill from the Technical Skills table
  create-variant     - Create a tailored resume copy for a job application
  create-cover-letter - Generate a formatted cover letter .docx
"""

import argparse
import copy
import os
import sys
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

RESUME_PATH = Path.home() / "Resume" / "MatthewDruhl.docx"
BACKUP_DIR = Path.home() / "Resume" / "archive"
APPLICATIONS_DIR = Path.home() / "Resume" / "applications"


def get_body_elements(doc):
    """Return list of (index, type, element) for all body children."""
    body = doc.element.body
    elements = []
    for i, child in enumerate(body):
        tag = etree.QName(child.tag).localname
        elements.append((i, tag, child))
    return elements


def get_paragraph_text(elem):
    """Extract full text from a paragraph element."""
    texts = [t.text for t in elem.findall(f".//{qn('w:t')}") if t.text]
    return "".join(texts)


KNOWN_SECTION_HEADERS = {
    "technical skills",
    "certifications",
    "professional experience",
    "additional relevant experience",
    "military service",
    "education",
}


def is_section_header(elem, after_first_known=False):
    """Check if a paragraph is a resume section header.

    Matches known headers by name. For custom/new sections, uses
    style-based detection (centered, bold, short) but only after
    the first known section header has been seen.
    """
    text = get_paragraph_text(elem).strip()
    if not text:
        return False

    # Known headers — always match
    if text.lower() in KNOWN_SECTION_HEADERS:
        return True

    # Style-based detection only after the header area
    if not after_first_known:
        return False

    jc = elem.find(f".//{qn('w:jc')}")
    if jc is None:
        return False
    if jc.get(qn("w:val")) != "center":
        return False

    bold = elem.find(f".//{qn('w:b')}")
    if bold is None:
        return False

    if len(text) > 40:
        return False
    if "Page Two" in text or "Page Three" in text:
        return False

    return True


def find_section_headers(doc):
    """Find all section headers with their body element indices."""
    body = doc.element.body
    headers = []
    found_first_known = False
    for i, child in enumerate(body):
        tag = etree.QName(child.tag).localname
        if tag == "p" and is_section_header(child, after_first_known=found_first_known):
            text = get_paragraph_text(child).strip()
            headers.append((i, text))
            if text.lower() in KNOWN_SECTION_HEADERS:
                found_first_known = True
    return headers


def cmd_read(args):
    """Display resume structure and content."""
    doc = Document(str(RESUME_PATH))
    body = doc.element.body
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    current_section = "Header"
    found_first_known = False
    print(f"{'='*60}")
    print(f"RESUME: {RESUME_PATH.name}")
    print(f"{'='*60}\n")

    for i, child in enumerate(body):
        tag = etree.QName(child.tag).localname

        if tag == "p":
            text = get_paragraph_text(child).strip()

            if is_section_header(child, after_first_known=found_first_known):
                if text.lower() in KNOWN_SECTION_HEADERS:
                    found_first_known = True
                current_section = text
                print(f"\n--- {text} {'—'*max(1, 45-len(text))}")
                continue

            if text:
                # Show bold portions in **bold**
                runs = child.findall(f".//{qn('w:r')}")
                formatted = []
                for run in runs:
                    run_text = "".join(
                        t.text for t in run.findall(qn("w:t")) if t.text
                    )
                    if not run_text:
                        continue
                    bold = run.find(f"{qn('w:rPr')}/{qn('w:b')}") is not None
                    if bold:
                        formatted.append(f"**{run_text}**")
                    else:
                        formatted.append(run_text)
                display = "".join(formatted)
                print(f"  {display}")

        elif tag == "tbl":
            rows = child.findall(f".//{qn('w:tr')}")
            for row in rows:
                cells = row.findall(f".//{qn('w:tc')}")
                cell_texts = []
                for cell in cells:
                    cell_text = "".join(
                        t.text
                        for t in cell.findall(f".//{qn('w:t')}")
                        if t.text
                    )
                    cell_texts.append(cell_text.strip())
                print(f"  | {' | '.join(cell_texts)} |")

    print(f"\n{'='*60}")


def cmd_page_count(args):
    """Convert resume to PDF via Word and count pages."""
    resume_path = Path(args.file) if args.file else RESUME_PATH

    if not resume_path.exists():
        print(f"Error: File not found at {resume_path}")
        sys.exit(1)

    try:
        from docx2pdf import convert
        from PyPDF2 import PdfReader
    except ImportError as e:
        print(f"Error: Missing dependency — {e}")
        print("Run with: uv run --with python-docx --with docx2pdf --with PyPDF2 python3 resume_tool.py page-count")
        sys.exit(1)

    with tempfile.TemporaryDirectory() as tmp_dir:
        pdf_path = Path(tmp_dir) / "resume.pdf"
        print(f"Converting {resume_path.name} to PDF via Word...")
        convert(str(resume_path), str(pdf_path))

        reader = PdfReader(str(pdf_path))
        page_count = len(reader.pages)

    print(f"\n{'='*40}")
    print(f"  File:  {resume_path.name}")
    print(f"  Pages: {page_count}")
    if page_count > 2:
        print(f"  ⚠  OVER 2-PAGE LIMIT by {page_count - 2} page(s)")
    elif page_count == 2:
        print(f"  ✓  At 2-page limit")
    else:
        print(f"  ✓  Under 2-page limit")
    print(f"{'='*40}")

    return page_count


def cmd_backup(args):
    """Create a timestamped backup of the resume."""
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"Druhl Matthew resume_{timestamp}.docx"
    backup_path = BACKUP_DIR / backup_name

    import shutil

    shutil.copy2(str(RESUME_PATH), str(backup_path))
    print(f"Backup created: {backup_path}")


def cmd_add_section(args):
    """Add a new section header after a specified existing section."""
    doc = Document(str(RESUME_PATH))
    body = doc.element.body
    headers = find_section_headers(doc)

    # Find the target section to insert after
    target_idx = None
    for idx, name in headers:
        if name.lower() == args.after.lower():
            target_idx = idx
            break

    if target_idx is None:
        available = [name for _, name in headers]
        print(f"Error: Section '{args.after}' not found.")
        print(f"Available sections: {', '.join(available)}")
        sys.exit(1)

    # Find the next section header to know where current section ends
    next_header_idx = None
    for idx, name in headers:
        if idx > target_idx:
            next_header_idx = idx
            break

    if next_header_idx is None:
        # Insert before the last element (sectPr)
        insert_before = body[-1]
    else:
        insert_before = body[next_header_idx]

    # Find the blank paragraph before the next section (if any)
    # We want to insert before that blank paragraph
    check_idx = list(body).index(insert_before)
    if check_idx > 0:
        prev = body[check_idx - 1]
        prev_tag = etree.QName(prev.tag).localname
        if prev_tag == "p" and not get_paragraph_text(prev).strip():
            insert_before = prev

    # Create section header matching existing style (copy from first found header)
    template_header = body[headers[0][0]]
    new_header = copy.deepcopy(template_header)

    # Clear runs and set new text
    for r in new_header.findall(qn("w:r")):
        new_header.remove(r)

    r = copy.deepcopy(template_header.findall(qn("w:r"))[0])
    for t in r.findall(qn("w:t")):
        r.remove(t)
    t = etree.SubElement(r, qn("w:t"))
    t.text = args.name
    new_header.append(r)

    # Create blank paragraphs
    blank_template = None
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p" and not get_paragraph_text(child).strip():
            blank_template = child
            break

    blank_before = copy.deepcopy(blank_template)
    blank_after = copy.deepcopy(blank_template)

    # Insert: blank, header, blank
    insert_before.addprevious(blank_before)
    blank_before.addnext(new_header)
    new_header.addnext(blank_after)

    doc.save(str(RESUME_PATH))
    print(f"Section '{args.name}' added after '{args.after}'.")


def cmd_add_entry(args):
    """Add an entry to an existing section."""
    doc = Document(str(RESUME_PATH))
    body = doc.element.body
    headers = find_section_headers(doc)

    # Find the target section
    target_idx = None
    for idx, name in headers:
        if name.lower() == args.section.lower():
            target_idx = idx
            break

    if target_idx is None:
        available = [name for _, name in headers]
        print(f"Error: Section '{args.section}' not found.")
        print(f"Available sections: {', '.join(available)}")
        sys.exit(1)

    # Find the next section header
    next_header_idx = None
    for idx, name in headers:
        if idx > target_idx:
            next_header_idx = idx
            break

    # Find the last content element in this section (before next section's blank line)
    if next_header_idx is None:
        # Last section — insert before sectPr
        insert_point = body[-1]
    else:
        insert_point = body[next_header_idx]
        # Walk back past blank paragraphs
        check_idx = list(body).index(insert_point)
        while check_idx > 0:
            prev = body[check_idx - 1]
            prev_tag = etree.QName(prev.tag).localname
            if prev_tag == "p" and not get_paragraph_text(prev).strip():
                insert_point = prev
                check_idx -= 1
            else:
                break

    # Create the entry paragraph
    # Use a clean blank paragraph as base
    blank_template = None
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p" and not get_paragraph_text(child).strip():
            blank_template = child
            break

    new_para = copy.deepcopy(blank_template)
    # Clear everything
    for child_elem in list(new_para):
        new_para.remove(child_elem)

    # Add bold run
    if args.bold_text:
        r1 = etree.SubElement(new_para, qn("w:r"))
        rPr1 = etree.SubElement(r1, qn("w:rPr"))
        etree.SubElement(rPr1, qn("w:b"))
        t1 = etree.SubElement(r1, qn("w:t"))
        t1.text = args.bold_text
        t1.set(qn("xml:space"), "preserve")

    # Add normal run
    if args.normal_text:
        r2 = etree.SubElement(new_para, qn("w:r"))
        t2 = etree.SubElement(r2, qn("w:t"))
        t2.text = args.normal_text
        t2.set(qn("xml:space"), "preserve")

    # Insert before the gap/next section
    insert_point.addprevious(new_para)

    doc.save(str(RESUME_PATH))
    full_text = (args.bold_text or "") + (args.normal_text or "")
    print(f"Entry added to '{args.section}': {full_text}")


def cmd_add_skill(args):
    """Add a skill to the Technical Skills table."""
    doc = Document(str(RESUME_PATH))

    if not doc.tables:
        print("Error: No tables found in the resume.")
        sys.exit(1)

    table = doc.tables[0]  # Technical Skills is the first table

    # Collect existing skills
    existing = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text:
                existing.append(text)

    if args.skill in existing:
        print(f"Skill '{args.skill}' already exists in the table.")
        return

    # Find first empty cell or add a new row
    added = False
    for row in table.rows:
        for cell in row.cells:
            if not cell.text.strip():
                cell.text = args.skill
                added = True
                break
        if added:
            break

    if not added:
        # Add a new row
        new_row = table.add_row()
        new_row.cells[0].text = args.skill

    doc.save(str(RESUME_PATH))
    print(f"Skill '{args.skill}' added to Technical Skills table.")


def cmd_remove_skill(args):
    """Remove a skill from the Technical Skills table and compact remaining cells."""
    file_path = Path(args.file) if args.file else RESUME_PATH
    doc = Document(str(file_path))

    if not doc.tables:
        print("Error: No tables found in the resume.")
        sys.exit(1)

    table = doc.tables[0]  # Technical Skills is the first table
    num_cols = len(table.columns)

    # Collect all skills, preserving cell formatting from first non-empty cell
    skills = []
    found = False
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text:
                if text.lower() == args.skill.lower():
                    found = True
                else:
                    skills.append(text)

    if not found:
        print(f"Error: Skill '{args.skill}' not found in the table.")
        existing = []
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    existing.append(t)
        print(f"Current skills: {', '.join(existing)}")
        sys.exit(1)

    # Clear all cells and refill with compacted list
    cell_idx = 0
    for row in table.rows:
        for cell in row.cells:
            # Clear existing text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.text = ""
            if cell_idx < len(skills):
                cell.paragraphs[0].runs[0].text = skills[cell_idx] if cell.paragraphs[0].runs else ""
                if not cell.paragraphs[0].runs:
                    cell.paragraphs[0].text = skills[cell_idx]
                cell_idx += 1

    # Remove empty trailing rows
    rows_to_remove = []
    for row in table.rows:
        if all(not cell.text.strip() for cell in row.cells):
            rows_to_remove.append(row)
    for row in rows_to_remove:
        row._element.getparent().remove(row._element)

    doc.save(str(file_path))
    print(f"Skill '{args.skill}' removed from Technical Skills table.")
    print(f"Remaining skills ({len(skills)}): {', '.join(skills)}")


def cmd_create_variant(args):
    """Copy base resume to an application-specific directory."""
    import shutil

    company = args.company
    title = args.title
    jobnumber = args.jobnumber or ""

    # Build directory name: title-jobnumber or just title
    dir_name = f"{title}-{jobnumber}" if jobnumber else title
    # Sanitize for filesystem
    dir_name = dir_name.replace(" ", "-").replace("/", "-")
    company_dir = company.replace(" ", "-").replace("/", "-")

    target_dir = APPLICATIONS_DIR / company_dir / dir_name
    target_dir.mkdir(parents=True, exist_ok=True)

    # Build filename: Resume-MATTHEW-DRUHL-{Company}.docx
    company_filename = company.replace(" ", "")
    target_file = target_dir / f"Resume-MATTHEW-DRUHL-{company_filename}.docx"

    if target_file.exists():
        print(f"Warning: Variant already exists at {target_file}")
        print("Overwriting with fresh copy from base resume.")

    shutil.copy2(str(RESUME_PATH), str(target_file))
    print(f"Variant created: {target_file}")
    print(f"Directory: {target_dir}")


def cmd_create_cover_letter(args):
    """Create a formatted cover letter .docx matching the Sparksoft style."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    target_dir = Path(args.target_dir)
    if not target_dir.exists():
        print(f"Error: Target directory does not exist: {target_dir}")
        sys.exit(1)

    # Read body content from file
    body_file = Path(args.body_file)
    if not body_file.exists():
        print(f"Error: Body file not found: {body_file}")
        sys.exit(1)

    body_paragraphs = body_file.read_text().strip().split("\n\n")

    company = args.company
    job_title = args.job_title
    date_str = args.date or datetime.now().strftime("%B %d, %Y")

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Date and company line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{date_str}\n{company}")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Salutation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Dear Hiring Manager,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Body paragraphs
    for para_text in body_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(para_text.strip())
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    # Signature block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Matthew Druhl")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tagline = "Empower Teams & Systems | Deliver Scalable Solutions & Seamless Collaboration"
    contact_lines = [
        tagline,
        "+1 603-377-1038",
        "matthewdruhl@gmail.com",
        "www.linkedin.com/in/matthew-druhl",
    ]
    for i, line in enumerate(contact_lines):
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if i < len(contact_lines) - 1:
            p.add_run("\n").font.size = Pt(11)

    # Save
    company_filename = company.replace(" ", "")
    output_path = target_dir / f"CoverLetter-MATTHEW-DRUHL-{company_filename}.docx"
    doc.save(str(output_path))
    print(f"Cover letter created: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Resume Editor Tool for MARVIN")
    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    # read
    subparsers.add_parser("read", help="Display resume structure and content")

    # page-count
    pc_parser = subparsers.add_parser("page-count", help="Check page count via Word PDF conversion")
    pc_parser.add_argument("--file", default="", help="Path to .docx file (defaults to active resume)")

    # backup
    subparsers.add_parser("backup", help="Create a timestamped backup")

    # add-section
    section_parser = subparsers.add_parser("add-section", help="Add a new section")
    section_parser.add_argument("--name", required=True, help="Section name")
    section_parser.add_argument(
        "--after", required=True, help="Insert after this section"
    )

    # add-entry
    entry_parser = subparsers.add_parser(
        "add-entry", help="Add entry to a section"
    )
    entry_parser.add_argument("--section", required=True, help="Target section name")
    entry_parser.add_argument("--bold-text", default="", help="Bold portion of entry")
    entry_parser.add_argument(
        "--normal-text", default="", help="Normal text portion of entry"
    )

    # add-skill
    skill_parser = subparsers.add_parser(
        "add-skill", help="Add skill to Technical Skills table"
    )
    skill_parser.add_argument("--skill", required=True, help="Skill name to add")

    # remove-skill
    rm_skill_parser = subparsers.add_parser(
        "remove-skill", help="Remove skill from Technical Skills table"
    )
    rm_skill_parser.add_argument("--skill", required=True, help="Skill name to remove")
    rm_skill_parser.add_argument("--file", default="", help="Path to .docx file (defaults to active resume)")

    # create-variant
    variant_parser = subparsers.add_parser(
        "create-variant", help="Create a tailored resume copy for a job application"
    )
    variant_parser.add_argument("--company", required=True, help="Company name")
    variant_parser.add_argument("--title", required=True, help="Job title")
    variant_parser.add_argument("--jobnumber", default="", help="Job number/ID (optional)")

    # create-cover-letter
    cl_parser = subparsers.add_parser(
        "create-cover-letter", help="Generate a formatted cover letter .docx"
    )
    cl_parser.add_argument("--target-dir", required=True, help="Directory to save the cover letter")
    cl_parser.add_argument("--company", required=True, help="Company name")
    cl_parser.add_argument("--job-title", required=True, help="Job title")
    cl_parser.add_argument("--date", default="", help="Date string (defaults to today)")
    cl_parser.add_argument("--body-file", required=True, help="Path to .txt file with letter body paragraphs")

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    if not RESUME_PATH.exists():
        print(f"Error: Resume not found at {RESUME_PATH}")
        sys.exit(1)

    commands = {
        "read": cmd_read,
        "page-count": cmd_page_count,
        "backup": cmd_backup,
        "add-section": cmd_add_section,
        "add-entry": cmd_add_entry,
        "add-skill": cmd_add_skill,
        "remove-skill": cmd_remove_skill,
        "create-variant": cmd_create_variant,
        "create-cover-letter": cmd_create_cover_letter,
    }

    commands[args.command](args)


if __name__ == "__main__":
    main()
