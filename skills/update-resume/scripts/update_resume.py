"""
Update Resume Tool for MARVIN

Scans cert PDFs, reads resume state, backs up, and applies updates
(add certifications, restructure Technical Skills table).

Usage: uv run --with python-docx,PyPDF2 python3 update_resume.py <command> [options]

Commands:
  scan-certs          - Extract text from ~/Resume/certs/*.pdf
  show-certs          - Show existing certifications on the resume
  show-skills         - Show current Technical Skills table
  backup              - Create timestamped backup
  add-certs           - Add certs from JSON (stdin) and restructure skills table
"""

import argparse
import copy
import json
import os
import stat
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Paths — configurable via environment variables (Issue #29, #51)
_RESUME_DIR = Path(os.environ.get("RESUME_DATA_DIR", str(Path.home() / "Resume")))
_default_docx = next(_RESUME_DIR.glob("*.docx"), _RESUME_DIR / "resume.docx")
RESUME_PATH = Path(os.environ.get("RESUME_DOCX_PATH", str(_default_docx)))
TEMPLATE_PATH = _RESUME_DIR / "original" / RESUME_PATH.name
BACKUP_DIR = _RESUME_DIR / "backup"
CERTS_DIR = _RESUME_DIR / "certs"

def safe_save(doc, path):
    """Save document, temporarily making file writable if needed."""
    path = Path(path)
    was_readonly = False
    if path.exists():
        mode = path.stat().st_mode
        if not (mode & stat.S_IWUSR):
            was_readonly = True
            path.chmod(mode | stat.S_IWUSR)
    doc.save(str(path))
    if was_readonly:
        path.chmod(path.stat().st_mode & ~(stat.S_IWUSR | stat.S_IWGRP | stat.S_IWOTH))


KNOWN_SECTION_HEADERS = {
    "technical skills",
    "certifications",
    "professional experience",
    "additional relevant experience",
    "military service",
    "education",
}


def get_paragraph_text(elem):
    """Extract full text from a paragraph element."""
    texts = [t.text for t in elem.findall(f".//{qn('w:t')}") if t.text]
    return "".join(texts)


def is_section_header(elem, after_first_known=False):
    """Check if a paragraph is a resume section header."""
    text = get_paragraph_text(elem).strip()
    if not text:
        return False
    if text.lower() in KNOWN_SECTION_HEADERS:
        return True
    if not after_first_known:
        return False
    jc = elem.find(f".//{qn('w:jc')}")
    if jc is None or jc.get(qn("w:val")) != "center":
        return False
    bold = elem.find(f".//{qn('w:b')}")
    if bold is None:
        return False
    if len(text) > 40 or "Page Two" in text or "Page Three" in text:
        return False
    return True


def find_section_headers(doc):
    """Find all section headers with their body element indices."""
    body = doc.element.body
    headers = []
    found_first_known = False
    for i, child in enumerate(body):
        tag = etree.QName(child.tag).localname
        if tag == "p" and is_section_header(child, after_first_known=found_first_known):
            text = get_paragraph_text(child).strip()
            headers.append((i, text))
            if text.lower() in KNOWN_SECTION_HEADERS:
                found_first_known = True
    return headers


def find_section_range(doc, section_name):
    """Find start and end indices for a section's content."""
    headers = find_section_headers(doc)
    body = doc.element.body
    target_idx = None
    for idx, name in headers:
        if name.lower() == section_name.lower():
            target_idx = idx
            break
    if target_idx is None:
        return None, None
    next_header_idx = None
    for idx, name in headers:
        if idx > target_idx:
            next_header_idx = idx
            break
    if next_header_idx is None:
        next_header_idx = len(list(body))
    return target_idx, next_header_idx


def cmd_scan_certs(args):
    """Extract text from all cert PDFs."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        print("Error: PyPDF2 not installed. Run with: uv run --with PyPDF2")
        sys.exit(1)

    pdf_files = sorted(CERTS_DIR.glob("*.pdf"))
    if not pdf_files:
        print("No PDF files found in ~/Resume/certs/")
        return

    for pdf_path in pdf_files:
        print(f"=== {pdf_path.name} ===")
        try:
            reader = PdfReader(str(pdf_path))
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    print(text)
        except Exception as e:
            print(f"Error reading {pdf_path.name}: {e}")
        print()


def cmd_show_certs(args):
    """Show existing certifications on the resume."""
    if not RESUME_PATH.exists():
        print(f"Error: Resume not found at {RESUME_PATH}")
        sys.exit(1)

    doc = Document(str(RESUME_PATH))
    body = doc.element.body
    start_idx, end_idx = find_section_range(doc, "certifications")

    if start_idx is None:
        print("No Certifications section found on resume.")
        return

    print("Existing certifications:")
    for i in range(start_idx + 1, end_idx):
        child = list(body)[i]
        tag = etree.QName(child.tag).localname
        if tag == "p":
            text = get_paragraph_text(child).strip()
            if text:
                print(f"  - {text}")


def cmd_show_skills(args):
    """Show current Technical Skills table."""
    if not RESUME_PATH.exists():
        print(f"Error: Resume not found at {RESUME_PATH}")
        sys.exit(1)

    doc = Document(str(RESUME_PATH))
    if not doc.tables:
        print("No tables found in resume.")
        return

    table = doc.tables[0]
    print(f"Technical Skills table: {len(table.rows)} rows x {len(table.columns)} columns")
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {i}: {' | '.join(cells)}")


def cmd_backup(args):
    """Create a timestamped backup of the resume."""
    if not RESUME_PATH.exists():
        print(f"Error: Resume not found at {RESUME_PATH}")
        sys.exit(1)

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_name = f"{RESUME_PATH.stem}_{timestamp}.docx"
    backup_path = BACKUP_DIR / backup_name

    import shutil
    shutil.copy2(str(RESUME_PATH), str(backup_path))
    print(f"Backup created: {backup_path}")


def cmd_add_certs(args):
    """Add certifications and restructure skills table.

    Reads JSON from stdin:
    {
        "certs": [
            {"name": "Cert Name", "issuer": "Org (Platform)", "date": "Mon YYYY"}
        ]
    }
    """
    if not RESUME_PATH.exists():
        print(f"Error: Resume not found at {RESUME_PATH}")
        sys.exit(1)

    # Read cert data from stdin
    try:
        data = json.load(sys.stdin)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON input: {e}")
        sys.exit(1)

    new_certs = data.get("certs", [])
    if not new_certs:
        print("No certs provided in JSON input.")
        return

    doc = Document(str(RESUME_PATH))
    body = doc.element.body
    headers = find_section_headers(doc)

    # Check for existing Certifications section
    cert_section_exists = any(
        name.lower() == "certifications" for _, name in headers
    )

    # Find existing cert texts to skip duplicates
    existing_cert_texts = set()
    if cert_section_exists:
        start_idx, end_idx = find_section_range(doc, "certifications")
        if start_idx is not None:
            for i in range(start_idx + 1, end_idx):
                child = list(body)[i]
                tag = etree.QName(child.tag).localname
                if tag == "p":
                    text = get_paragraph_text(child).strip().lower()
                    if text:
                        existing_cert_texts.add(text)

    # Filter out certs already on resume
    certs_to_add = []
    for cert in new_certs:
        # Check if cert name already appears in any existing cert text
        name_lower = cert["name"].lower()
        if any(name_lower in existing for existing in existing_cert_texts):
            print(f"Skipping (already on resume): {cert['name']}")
        else:
            certs_to_add.append(cert)

    if not certs_to_add:
        print("All certs already on resume. Nothing to add.")
        # Still restructure skills table
        _restructure_skills_table(doc)
        safe_save(doc, RESUME_PATH)
        print("Technical Skills table restructured.")
        return

    # Find insertion point: after Technical Skills, before Professional Experience
    tech_skills_idx = None
    prof_exp_idx = None
    for idx, name in headers:
        if name.lower() == "technical skills":
            tech_skills_idx = idx
        if name.lower() == "professional experience":
            prof_exp_idx = idx

    if tech_skills_idx is None or prof_exp_idx is None:
        print("Error: Could not find Technical Skills or Professional Experience section.")
        sys.exit(1)

    if not cert_section_exists:
        # Create Certifications section between Technical Skills and Professional Experience
        # Find the element just before Professional Experience header
        prof_exp_elem = list(body)[prof_exp_idx]

        # Walk back to find blank paragraph before Professional Experience
        insert_before = prof_exp_elem
        check_idx = prof_exp_idx
        while check_idx > 0:
            prev = list(body)[check_idx - 1]
            prev_tag = etree.QName(prev.tag).localname
            if prev_tag == "p" and not get_paragraph_text(prev).strip():
                insert_before = prev
                check_idx -= 1
            else:
                break

        # Create section header matching existing style
        # Find a section header to use as template
        template_header = None
        for idx, name in headers:
            if name.lower() in KNOWN_SECTION_HEADERS:
                template_header = list(body)[idx]
                break

        if template_header is None:
            print("Error: No section header template found.")
            sys.exit(1)

        # Create the "Certifications" header
        new_header = copy.deepcopy(template_header)
        for r in new_header.findall(qn("w:r")):
            new_header.remove(r)
        r = copy.deepcopy(template_header.findall(qn("w:r"))[0])
        for t in r.findall(qn("w:t")):
            r.remove(t)
        t = etree.SubElement(r, qn("w:t"))
        t.text = "Certifications"
        new_header.append(r)

        # Find a blank paragraph template
        blank_template = None
        for child in body:
            tag = etree.QName(child.tag).localname
            if tag == "p" and not get_paragraph_text(child).strip():
                blank_template = child
                break

        # Insert: blank, header (no trailing blank — keep compact for 2-page limit)
        blank_before = copy.deepcopy(blank_template)

        insert_before.addprevious(blank_before)
        blank_before.addnext(new_header)

        print("Created Certifications section.")

        # Refresh headers after modification
        headers = find_section_headers(doc)

    # Now add cert entries to the Certifications section
    cert_start, cert_end = find_section_range(doc, "certifications")
    if cert_start is None:
        print("Error: Certifications section not found after creation.")
        sys.exit(1)

    # Find a content paragraph to use as template for font/style
    content_template = None
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p":
            text = get_paragraph_text(child).strip()
            if text and not is_section_header(child, after_first_known=True):
                runs = child.findall(qn("w:r"))
                if runs:
                    content_template = child
                    break

    # Collect existing cert data from section
    all_certs = []
    existing_para_elems = []
    for i in range(cert_start + 1, cert_end):
        child = list(body)[i]
        tag = etree.QName(child.tag).localname
        if tag == "p":
            text = get_paragraph_text(child).strip()
            if text:
                existing_para_elems.append(child)
                # Parse cert info from text: "Name, Issuer, Mon YYYY"
                parts = text.rsplit(", ", 2)
                if len(parts) >= 3:
                    all_certs.append({"name": parts[0], "issuer": parts[1], "date": parts[2], "source": "existing"})

    # Add new certs
    for cert in certs_to_add:
        all_certs.append({**cert, "source": "new"})

    # Sort all certs by date descending (newest first)
    all_certs.sort(key=lambda c: _parse_cert_date(c["date"]), reverse=True)

    # Remove existing cert paragraphs from section
    for elem in existing_para_elems:
        body.remove(elem)

    # Re-find section range after removals
    cert_start, cert_end = find_section_range(doc, "certifications")

    # Insert directly after the Certifications header (no blank line between)
    insert_after_elem = list(body)[cert_start]

    # Insert all certs in reverse chronological order
    for cert in reversed(all_certs):
        new_para = _create_cert_paragraph(cert, content_template, body)
        insert_after_elem.addnext(new_para)
        label = "Added" if cert.get("source") == "new" else "Kept"
        print(f"{label}: {cert['name']}, {cert['issuer']}, {cert['date']}")

    # Remove any blank paragraphs between last cert and next section header
    # to keep the resume compact for the 2-page limit
    _remove_trailing_blanks(doc, "certifications")

    # Restructure Technical Skills table
    _restructure_skills_table(doc)

    # Fix page breaks to prevent blank pages
    _fix_page_breaks(doc)

    safe_save(doc, RESUME_PATH)
    print(f"\nResume saved to: {RESUME_PATH}")
    print("IMPORTANT: Open the document to verify it stays within 2 pages.")


def _parse_cert_date(date_str):
    """Parse 'Mon YYYY' date string for sorting. Returns datetime."""
    try:
        return datetime.strptime(date_str.strip(), "%b %Y")
    except ValueError:
        return datetime.min


def _remove_trailing_blanks(doc, section_name):
    """Remove blank paragraphs between a section's last content and the next section header."""
    body = doc.element.body
    start_idx, end_idx = find_section_range(doc, section_name)
    if start_idx is None:
        return
    children = list(body)
    # Walk backwards from end of section, removing blank paragraphs
    for i in range(end_idx - 1, start_idx, -1):
        child = children[i]
        tag = etree.QName(child.tag).localname
        if tag == "p" and not get_paragraph_text(child).strip():
            body.remove(child)
        else:
            break


def _create_cert_paragraph(cert, content_template, body):
    """Create a paragraph for a certification entry.

    Format: **Cert Name**, Issuer, Date
    """
    # Use blank paragraph as base
    blank_template = None
    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "p" and not get_paragraph_text(child).strip():
            blank_template = child
            break

    new_para = copy.deepcopy(blank_template)
    # Clear all children
    for child_elem in list(new_para):
        new_para.remove(child_elem)

    # Copy paragraph properties from content template if available
    if content_template is not None:
        pPr = content_template.find(qn("w:pPr"))
        if pPr is not None:
            new_pPr = copy.deepcopy(pPr)
            # Remove centering if present (certs should be left-aligned)
            jc = new_pPr.find(qn("w:jc"))
            if jc is not None:
                new_pPr.remove(jc)
            new_para.insert(0, new_pPr)

    # Get font properties from content template
    rPr_template = None
    if content_template is not None:
        first_run = content_template.find(qn("w:r"))
        if first_run is not None:
            rPr_template = first_run.find(qn("w:rPr"))

    # Bold run: cert name
    r1 = etree.SubElement(new_para, qn("w:r"))
    rPr1 = etree.SubElement(r1, qn("w:rPr"))
    etree.SubElement(rPr1, qn("w:b"))
    # Copy font info from template
    if rPr_template is not None:
        for child in rPr_template:
            tag = etree.QName(child.tag).localname
            if tag in ("rFonts", "sz", "szCs"):
                rPr1.append(copy.deepcopy(child))
    t1 = etree.SubElement(r1, qn("w:t"))
    t1.text = cert["name"]
    t1.set(qn("xml:space"), "preserve")

    # Normal run: , Issuer, Date
    r2 = etree.SubElement(new_para, qn("w:r"))
    if rPr_template is not None:
        rPr2 = copy.deepcopy(rPr_template)
        # Remove bold if present in template
        for b in rPr2.findall(qn("w:b")):
            rPr2.remove(b)
        r2.insert(0, rPr2)
    t2 = etree.SubElement(r2, qn("w:t"))
    t2.text = f", {cert['issuer']}, {cert['date']}"
    t2.set(qn("xml:space"), "preserve")

    return new_para


def _restructure_skills_table(doc):
    """Restructure Technical Skills table to 4 columns, alphabetized."""
    if not doc.tables:
        print("No tables found — skipping skills restructure.")
        return

    table = doc.tables[0]

    # Collect all skill items
    skills = []
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text:
                skills.append(text)

    if not skills:
        print("No skills found in table.")
        return

    # Alphabetize
    skills.sort(key=str.lower)

    # Calculate grid: 4 columns
    num_cols = 4
    num_rows = (len(skills) + num_cols - 1) // num_cols

    # Get the table XML element
    tbl = table._tbl

    # Get table properties
    tblPr = tbl.find(qn("w:tblPr"))
    tblGrid = tbl.find(qn("w:tblGrid"))

    # Remove all existing rows
    for tr in tbl.findall(qn("w:tr")):
        tbl.remove(tr)

    # Remove existing grid and create new one with 4 columns
    if tblGrid is not None:
        tbl.remove(tblGrid)

    new_grid = etree.SubElement(tbl, qn("w:tblGrid"))
    # Calculate column width (assuming full page width ~9360 twips for 6.5" content)
    col_width = 9360 // num_cols
    for _ in range(num_cols):
        gridCol = etree.SubElement(new_grid, qn("w:gridCol"))
        gridCol.set(qn("w:w"), str(col_width))

    # Move grid after tblPr
    if tblPr is not None:
        tblPr.addnext(new_grid)

    # Create new rows
    for row_idx in range(num_rows):
        tr = etree.SubElement(tbl, qn("w:tr"))
        for col_idx in range(num_cols):
            skill_idx = row_idx * num_cols + col_idx
            tc = etree.SubElement(tr, qn("w:tc"))

            # Set cell width
            tcPr = etree.SubElement(tc, qn("w:tcPr"))
            tcW = etree.SubElement(tcPr, qn("w:tcW"))
            tcW.set(qn("w:w"), str(col_width))
            tcW.set(qn("w:type"), "dxa")

            # Add paragraph with skill text
            p = etree.SubElement(tc, qn("w:p"))
            if skill_idx < len(skills):
                r = etree.SubElement(p, qn("w:r"))
                t = etree.SubElement(r, qn("w:t"))
                t.text = skills[skill_idx]

    print(f"Technical Skills table restructured: {num_rows} rows x {num_cols} columns")
    print(f"Skills ({len(skills)}): {', '.join(skills)}")


def _fix_page_breaks(doc):
    """Replace explicit page break paragraphs with pageBreakBefore on Page Two header.

    Prevents blank pages when content is added before the page break.
    """
    body = doc.element.body
    page_break_elems = []
    page_two_elem = None

    for child in body:
        tag = etree.QName(child.tag).localname
        if tag != "p":
            continue
        text = get_paragraph_text(child).strip()
        br = child.find(f".//{qn('w:br')}[@{qn('w:type')}='page']")
        if br is not None and not text:
            page_break_elems.append(child)
        if "Page Two" in text:
            page_two_elem = child

    for elem in page_break_elems:
        body.remove(elem)
        print("Removed explicit page break paragraph")

    if page_two_elem is not None:
        pPr = page_two_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = etree.SubElement(page_two_elem, qn("w:pPr"))
            page_two_elem.insert(0, pPr)
        pbf = pPr.find(qn("w:pageBreakBefore"))
        if pbf is None:
            etree.SubElement(pPr, qn("w:pageBreakBefore"))
            print("Added pageBreakBefore to Page Two header")


def main():
    parser = argparse.ArgumentParser(description="Update Resume Tool for MARVIN")
    subparsers = parser.add_subparsers(dest="command", help="Available commands")

    subparsers.add_parser("scan-certs", help="Extract text from cert PDFs")
    subparsers.add_parser("show-certs", help="Show existing certifications on resume")
    subparsers.add_parser("show-skills", help="Show current Technical Skills table")
    subparsers.add_parser("backup", help="Create timestamped backup")
    subparsers.add_parser(
        "add-certs",
        help="Add certs from JSON (stdin) and restructure skills table",
    )

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        sys.exit(1)

    commands = {
        "scan-certs": cmd_scan_certs,
        "show-certs": cmd_show_certs,
        "show-skills": cmd_show_skills,
        "backup": cmd_backup,
        "add-certs": cmd_add_certs,
    }

    commands[args.command](args)


if __name__ == "__main__":
    main()
